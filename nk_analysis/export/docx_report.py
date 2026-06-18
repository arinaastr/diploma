import os
import tempfile
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from nk_analysis.core.chart import draw_scatter, save_chart_to_temp, cleanup_temp
from nk_analysis.core.math_engine import get_beton_class, beton_class_index

def _heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    return p

def _add_table(doc, df, cols=None):
    if df is None or len(df) == 0:
        doc.add_paragraph("Нет данных.")
        return
    if cols is None:
        cols = list(df.columns)
    cols = [c for c in cols if c in df.columns]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = str(c)
        hdr[i].paragraphs[0].runs[0].bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            val = row[c]
            cells[i].text = "—" if (val is None or (isinstance(val, float) and np.isnan(val))) else str(val)
    doc.add_paragraph("")

def generate_docx(state, meta, include_chart=True,
                  include_stvol=True, include_ne=True):
    doc = Document()

    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    title = doc.add_paragraph("ПРОТОКОЛ")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    sub = doc.add_paragraph(
        "испытания прочности бетона методом ультразвукового прозвучивания"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph(f"№ {meta.get('num', '—')}  от  {meta.get('date', '—')}")

    _heading(doc, "1. Общие сведения")
    info = [
        ("Объект",              meta.get("obj",    "—")),
        ("Адрес",               meta.get("addr",   "—")),
        ("Период обследования", meta.get("period", "—")),
        ("Приборы",             meta.get("dev",    "—")),
        ("НТД",                 meta.get("ntd",    "—")),
        ("Возраст бетона, сут", str(meta.get("age", "—"))),
        ("Проектный класс бетона", meta.get("proj_cls", "—")),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    doc.add_paragraph("")

    def _write_cal_block(doc, cal, section_title, chart_title, iter_label, data_df, data_cols, include_chart, data_heading="Данные по элементам"):
        _heading(doc, section_title)
        if cal:
            a0, a1 = cal["a0"], cal["a1"]
            b_str = f"+ {a0:.3f}" if a0 >= 0 else f"- {abs(a0):.3f}"
            doc.add_paragraph(
                f"Градуировочная зависимость: R = {a1:.5f}·V {b_str}"
            )
            sr = cal.get("sr", float("nan"))
            sr_str = "—" if (sr != sr) else f"{sr:.4f}"
            doc.add_paragraph(
                f"a = {a1:.5f},  b = {a0:.3f},  S = {cal['S_T']:.3f} МПа,  "
                f"r = {cal['r']:.4f},  S/R = {sr_str}"
            )
            if not cal.get("valid", True):
                p_warn = doc.add_paragraph(
                    "ВНИМАНИЕ: Градуировочная зависимость НЕ ДОПУСТИМА "
                    "(коэффициент корреляции r < 0.7 или S/R > 0.15). "
                    "Контроль и оценка прочности по данной градуировочной зависимости не допускаются (ГОСТ 17624)."
                )
                p_warn.runs[0].bold = True
                p_warn.runs[0].font.color.rgb = RGBColor(0x8B, 0x1A, 0x1A)

            if cal["iters"]:
                _heading(doc, f"Итерации отбраковки ({iter_label})", level=2)
                _add_table(doc, pd.DataFrame(cal["iters"]))
            if include_chart:
                p_chart_title = doc.add_paragraph(chart_title)
                p_chart_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_chart_title.runs[0].bold = True
                fig, ax = plt.subplots(figsize=(14, 5))
                draw_scatter(ax, fig, cal, title="")
                try:
                    tmp = save_chart_to_temp(fig)
                finally:
                    plt.close(fig)
                try:
                    doc.add_picture(tmp, width=Cm(14))
                finally:
                    cleanup_temp(tmp)
                doc.add_paragraph("")
        if data_df is not None and len(data_df):
            _heading(doc, data_heading, level=2)
            cols = [c for c in data_cols if c in data_df.columns]
            _add_table(doc, data_df, cols)

    if include_stvol:
        dc_s = state.get("dc_stvol", pd.DataFrame())
        cols_s = ["Горизонт", "Сторона", "V", "f_МО", "f_расч МПа", "Класс", "Статус"]
        _write_cal_block(doc, state.get("cal_s"), "2. Результаты по стволу",
                         "Градуировочная зависимость — Ствол", "ствол",
                         dc_s if len(dc_s) else None, cols_s, include_chart,
                         data_heading="Данные по стволу")

    if include_ne:
        dc_n = state.get("dc_ne", pd.DataFrame())
        cols_n = ["Участок", "V", "f_МО", "f_расч МПа", "Класс", "Статус"]
        _write_cal_block(doc, state.get("cal_n"), "3. Результаты по конструкциям (не ствол)",
                         "Градуировочная зависимость — Конструкции", "конструкции",
                         dc_n if len(dc_n) else None, cols_n, include_chart,
                         data_heading="Данные по конструкциям")

    _heading(doc, "4. Заключение")

    frames = []
    dc_s = state.get("dc_stvol", pd.DataFrame())
    dc_n = state.get("dc_ne",    pd.DataFrame())
    if len(dc_s) and "f_расч МПа" in dc_s.columns:
        frames.append(dc_s["f_расч МПа"].dropna())
    if len(dc_n) and "f_расч МПа" in dc_n.columns:
        frames.append(dc_n["f_расч МПа"].dropna())

    fact_cls = "—"
    rm_str   = "—"
    if frames:
        all_f = pd.concat(frames)
        if len(all_f):
            rm = all_f.mean()
            rm_str   = f"{rm:.1f}"
            fact_cls = get_beton_class(rm)

    proj_cls = meta.get("proj_cls", "—")
    doc.add_paragraph(f"Средняя расчётная прочность: {rm_str} МПа")
    doc.add_paragraph(f"Фактический класс бетона:   {fact_cls}")
    doc.add_paragraph(f"Проектный класс бетона:     {proj_cls}")

    if fact_cls != "—":
        if proj_cls and proj_cls not in ("—", "", "— не указан —"):
            fi = beton_class_index(fact_cls)
            pi = beton_class_index(proj_cls)
            if fi == -1 or pi == -1:
                verdict = f"Фактический класс бетона: {fact_cls}. Сравнение с проектным ({proj_cls}) невозможно — неизвестный класс."
            elif fi >= pi:
                verdict = f"Фактический класс бетона СООТВЕТСТВУЕТ проектному ({fact_cls} ≥ {proj_cls})."
            else:
                verdict = f"Фактический класс бетона НЕ СООТВЕТСТВУЕТ проектному ({fact_cls} < {proj_cls})."
        else:
            verdict = f"По результатам обследования бетон имеет класс {fact_cls} (средняя прочность {rm_str} МПа)."
        p = doc.add_paragraph(verdict)
        p.runs[0].bold = True

    _heading(doc, "5. Исполнители")
    sign_tbl = doc.add_table(rows=2, cols=2)
    sign_tbl.style = "Table Grid"
    e1f = meta.get('e1f','') or '—'
    e1p = meta.get('e1p','') or ''
    e2f = meta.get('e2f','') or '—'
    e2p = meta.get('e2p','') or ''
    sign_tbl.rows[0].cells[0].text = f"Исп. 1: {e1f}" + (f", {e1p}" if e1p else "")
    sign_tbl.rows[1].cells[0].text = f"Исп. 2: {e2f}" + (f", {e2p}" if e2p else "")
    sign_tbl.rows[0].cells[1].text = "Подпись: ____________"
    sign_tbl.rows[1].cells[1].text = "Подпись: ____________"

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path
